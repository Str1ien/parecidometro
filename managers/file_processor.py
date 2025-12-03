"""
File Processor Module

This module handles the processing of different file types to extract
content suitable for similarity hashing.

Classes:
    FileProcessor: Main class for file type detection and content extraction

Supported File Types:
    - PDF documents (text extraction)
    - Microsoft Word DOCX (text extraction)
    - Binary executables (PE/ELF)
    - Generic files (raw content)

Dependencies:
    - python-magic: File type detection
    - python-docx: Word document processing
    - PyMuPDF (fitz): PDF text extraction
"""

import magic
from docx import Document
import fitz  # PyMuPDF
import logging

logger = logging.getLogger(__name__)


class FileProcessor:
    """
    Processes different file types and extracts content for TLSH/ssdeep hashing.

    This class detects file types and extracts appropriate content for similarity
    hashing. For documents (PDF, DOCX), it extracts text. For binaries and
    other files, it uses raw content.

    Attributes:
        file_data (bytes): Raw file content
        filename (str): Original filename
        file_type (str): Detected MIME type
    """

    def __init__(self, file_data, filename):
        """
        Initialize the FileProcessor.

        Args:
            file_data (bytes): Raw file content
            filename (str): Original filename (used for logging)
        """
        self.file_data = file_data
        self.filename = filename
        self.file_type = self._detect_file_type()

    def _detect_file_type(self):
        """
        Detect file type using libmagic.

        Uses python-magic library to detect the MIME type of the file
        based on its content (not just the extension).

        Returns:
            str: MIME type (e.g., "application/pdf", "text/plain")

        Example:
            >>> processor = FileProcessor(data, "file.pdf")
            >>> mime_type = processor._detect_file_type()
            >>> print(mime_type)  # "application/pdf"
        """
        mime = magic.Magic(mime=True)
        detected_type = mime.from_buffer(self.file_data)
        logger.debug(f"File type detected for {self.filename}: {detected_type}")
        return detected_type

    def get_file_type(self):
        """
        Get the detected MIME type.

        Returns:
            str: MIME type of the file
        """
        return self.file_type

    def process(self):
        """
        Process the file according to its type.

        This method routes the file to the appropriate processing function
        based on its detected MIME type:
        - PDF: Extract text
        - DOCX: Extract text
        - Executables: Return raw content
        - Others: Return raw content

        Returns:
            tuple: (success: bool, content: bytes | error_message: str)
                - If successful: (True, processed_content)
                - If failed: (False, "Error message")

        Example:
            >>> processor = FileProcessor(data, "document.pdf")
            >>> success, content = processor.process()
            >>> if success:
            ...     print(f"Extracted {len(content)} bytes")
        """
        logger.info(f"Processing file: {self.filename} (type: {self.file_type})")

        # PDF
        if "pdf" in self.file_type:
            return self._process_pdf()

        # Word
        elif "word" in self.file_type or "officedocument" in self.file_type:
            return self._process_docx()

        # Executables and other files: return raw content
        elif (
            "x-executable" in self.file_type
            or "x-dosexec" in self.file_type
            or "x-sharedlib" in self.file_type
            or "elf" in self.file_type.lower()
        ):
            return self._process_binary()

        # Any other file type: raw content
        else:
            return self._process_generic()

    def _process_binary(self):
        """
        Process binary executables (PE/ELF) - return raw content.

        Binary files (executables, libraries) are used as-is without
        any text extraction since their structure is important for
        similarity detection.

        Returns:
            tuple: (success: bool, content: bytes | error_message: str)

        Example:
            >>> processor = FileProcessor(exe_data, "malware.exe")
            >>> success, content = processor._process_binary()
            >>> if success:
            ...     print(f"Binary size: {len(content)} bytes")
        """
        size = len(self.file_data)
        logger.debug(f"Processing binary: {self.filename} ({size} bytes)")

        if size < 50:
            logger.warning(f"Binary file too small: {self.filename} ({size} bytes)")
            return False, "Binary file too small (min 50 bytes for TLSH)"

        logger.info(f"Binary processed successfully: {self.filename}")
        return True, self.file_data

    def _process_generic(self):
        """
        Process any generic file - return raw content.

        Generic files (text, scripts, unknown types) are processed using
        their raw content without any special handling.

        Returns:
            tuple: (success: bool, content: bytes | error_message: str)

        Example:
            >>> processor = FileProcessor(text_data, "script.py")
            >>> success, content = processor._process_generic()
        """
        size = len(self.file_data)
        logger.debug(f"Processing generic file: {self.filename} ({size} bytes)")

        if size < 50:
            logger.warning(f"Generic file too small: {self.filename} ({size} bytes)")
            return False, f"File too small (min 50 bytes for TLSH, got {size})"

        logger.info(f"Generic file processed successfully: {self.filename}")
        return True, self.file_data

    def _process_pdf(self):
        """
        Extract text from PDF documents.

        Uses PyMuPDF (fitz) to extract text from all pages of a PDF document.
        The extracted text is then used for similarity hashing, which allows
        detection of similar documents even if they have different formatting.

        Returns:
            tuple: (success: bool, content: bytes | error_message: str)
                - If successful: (True, utf8_encoded_text)
                - If failed: (False, "Error description")

        Notes:
            - Scanned PDFs without OCR will have no extractable text
            - Password-protected PDFs will fail
            - Corrupted PDFs will return an error

        Example:
            >>> processor = FileProcessor(pdf_data, "document.pdf")
            >>> success, text = processor._process_pdf()
            >>> if success:
            ...     print(f"Extracted text: {text.decode('utf-8')[:100]}...")
        """
        import io

        try:
            logger.debug(f"Extracting text from PDF: {self.filename}")
            doc = fitz.open(stream=self.file_data, filetype="pdf")
            text = ""
            page_count = len(doc)

            for page in doc:
                text += page.get_text()
            doc.close()

            text_length = len(text.strip())

            if not text or text_length < 1:
                logger.warning(
                    f"PDF has no extractable text: {self.filename} ({page_count} pages)"
                )
                return (
                    False,
                    f"PDF has no extractable text ({page_count} pages scanned)",
                )

            logger.info(
                f"PDF processed successfully: {self.filename} ({page_count} pages, {text_length} chars extracted)"
            )
            return True, text.encode("utf-8")

        except Exception as e:
            logger.error(f"Error extracting PDF {self.filename}: {e}", exc_info=True)
            return (
                False,
                f"PDF extraction failed: {str(e)} - file may be corrupted or password-protected",
            )

    def _process_docx(self):
        """
        Extract text from Microsoft Word DOCX documents.

        Uses python-docx to extract text from all paragraphs in a DOCX file.
        The extracted text is used for similarity hashing.

        Returns:
            tuple: (success: bool, content: bytes | error_message: str)
                - If successful: (True, utf8_encoded_text)
                - If failed: (False, "Error description")

        Notes:
            - Only supports .docx (Office 2007+), not old .doc format
            - Extracts text from paragraphs, not from headers/footers/tables
            - Password-protected files will fail

        Example:
            >>> processor = FileProcessor(docx_data, "report.docx")
            >>> success, text = processor._process_docx()
            >>> if success:
            ...     print(f"Extracted text: {text.decode('utf-8')[:100]}...")
        """
        import io

        try:
            logger.debug(f"Extracting text from DOCX: {self.filename}")
            doc = Document(io.BytesIO(self.file_data))
            text = "\n".join([para.text for para in doc.paragraphs])

            text_length = len(text.strip())
            para_count = len(doc.paragraphs)

            if not text or text_length < 1:
                logger.warning(
                    f"DOCX has no extractable text: {self.filename} ({para_count} paragraphs)"
                )
                return (
                    False,
                    f"DOCX has no extractable text ({para_count} paragraphs found but all empty)",
                )

            logger.info(
                f"DOCX processed successfully: {self.filename} ({para_count} paragraphs, {text_length} chars extracted)"
            )
            return True, text.encode("utf-8")

        except Exception as e:
            logger.error(f"Error extracting DOCX {self.filename}: {e}", exc_info=True)
            return (
                False,
                f"DOCX extraction failed: {str(e)} - file may be corrupted or unsupported format",
            )
